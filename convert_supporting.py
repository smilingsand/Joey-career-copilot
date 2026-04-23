import json
import os
import re
import sys
import glob # [新增] 用于处理通配符

# 尝试导入 python-docx 以支持 word 文档
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn 
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("提示: 未检测到 'python-docx' 库。如需处理 .docx 文件，请运行: pip install python-docx")

def clean_json_string(json_str):
    json_str = re.sub(r',\s*([\]}])', r'\1', json_str)
    return json_str

def read_file_content(filepath):
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()

    if ext == '.docx':
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx 库来读取 .docx 文件")
        doc = Document(filepath)
        return '\n'.join([p.text for p in doc.paragraphs])
    else:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()

def write_file_content(filepath, content):
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()

    if ext == '.docx':
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx 库来写入 .docx 文件")
        
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'DengXian'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'DengXian') 
        style.paragraph_format.line_spacing = 1.16
        style.paragraph_format.space_after = Pt(8)
        
        for line in content.split('\n'):
            stripped_line = line.strip()
            
            if stripped_line.startswith("- "):
                p = doc.add_paragraph(style='List Bullet')
                text_content = stripped_line[2:]
            else:
                p = doc.add_paragraph()
                text_content = line
            
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            parts = re.split(r'(\*\*.*?\*\*)', text_content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = 'DengXian'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DengXian')
                else:
                    run = p.add_run(part)
                    run.font.name = 'DengXian'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'DengXian')

        doc.save(filepath)
    else:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

def clean_list_item(line):
    line = re.sub(r'^[\s\-\*]+', '', line)
    line = re.sub(r'^\d+\.\s*', '', line)
    return line.strip()

def get_output_base_path(input_path, prefix_from, prefix_to):
    dirname = os.path.dirname(input_path)
    filename = os.path.basename(input_path)
    
    if filename.startswith(prefix_from):
        out_filename = filename.replace(prefix_from, prefix_to, 1)
    else:
        out_filename = prefix_to + filename
        
    out_filename_no_ext = os.path.splitext(out_filename)[0]
    
    if dirname:
        return os.path.join(dirname, out_filename_no_ext)
    else:
        return out_filename_no_ext

def _parse_filename_metadata(filename):
    base = os.path.splitext(filename)[0]
    if base.startswith("Supporting_"):
        base = base.replace("Supporting_", "", 1)
    parts = base.split('_')
    
    metadata = {"title": "Unknown", "company": "Unknown"}
    if len(parts) >= 1: metadata["title"] = parts[0]
    if len(parts) >= 2: metadata["company"] = parts[1]
    return metadata

def support_to_debug(input_path):
    print(f"正在处理: {os.path.basename(input_path)} ...")
    try:
        content = read_file_content(input_path)
        metadata = _parse_filename_metadata(os.path.basename(input_path))
    
        debug_data = {
            "original_file": os.path.basename(input_path),
            "metadata": metadata,
            "requirements": [],
            "requirements_evidence": {},
            "findings": {},
            "validation_history": [],
            "cover_letter_content": "",
            "resume_summary_content": ""
        }
    
        skill_blocks = re.split(r'(?:\*\*)?\[Key skills and experience\](?:\*\*)?', content)
    
        for block in skill_blocks:
            if not block.strip(): continue
    
            lines = block.strip().split('\n')
            skill_name = lines[0].strip()
            if not skill_name: continue
    
            debug_data["requirements"].append(skill_name)
            debug_data["requirements_evidence"][skill_name] = []
            debug_data["findings"][skill_name] = {"materials": []}
    
            req_match = re.search(r'(?:\*\*)?\[Requirements\](?:\*\*)?(.*?)(?=(?:\*\*)?\[Supporting\]|$)', block, re.DOTALL)
            if req_match:
                req_lines = req_match.group(1).strip().split('\n')
                clean_reqs = [clean_list_item(l) for l in req_lines if l.strip()]
                debug_data["requirements_evidence"][skill_name] = clean_reqs
    
            sup_match = re.search(r'(?:\*\*)?\[Supporting\](?:\*\*)?(.*)', block, re.DOTALL)
            if sup_match:
                sup_lines = sup_match.group(1).strip().split('\n')
                clean_sups = [clean_list_item(l) for l in sup_lines if l.strip()]
                debug_data["findings"][skill_name]["materials"] = clean_sups
    
        out_path_no_ext = get_output_base_path(input_path, "Supporting_", "DEBUG_")
        out_filename = out_path_no_ext + ".json"
        
        with open(out_filename, 'w', encoding='utf-8') as f:
            json.dump(debug_data, f, indent=2, ensure_ascii=False)
        
        print(f"  -> 生成: {os.path.basename(out_filename)}")
        return True
    except Exception as e:
        print(f"  -> 失败: {e}")
        return False

def debug_to_support(input_path, output_format="txt"):
    print(f"正在处理: {os.path.basename(input_path)} ...")
    try:
        json_str = read_file_content(input_path)
        
        try:
            data = json.loads(json_str)
        except json.JSONDecodeError:
            json_str = clean_json_string(json_str)
            data = json.loads(json_str)
    
        output_lines = []
        skills = data.get("requirements", [])
        
        for skill in skills:
            output_lines.append(f"**[Key skills and experience]** {skill}")
            
            output_lines.append("**[Requirements]**")
            evidences = data.get("requirements_evidence", {}).get(skill, [])
            if isinstance(evidences, list):
                for ev in evidences:
                    output_lines.append(f"- {ev}")
            
            output_lines.append("**[Supporting]**")
            findings = data.get("findings", {}).get(skill, {})
            materials = findings.get("materials", []) if findings else []
            if isinstance(materials, list):
                for mat in materials:
                    output_lines.append(f"- {mat}")
            
            output_lines.append("") 
    
        out_path_no_ext = get_output_base_path(input_path, "DEBUG_", "Supporting_")
        out_filename = f"{out_path_no_ext}.{output_format}"
    
        write_file_content(out_filename, "\n".join(output_lines))
        
        print(f"  -> 生成: {os.path.basename(out_filename)}")
        return True
    except Exception as e:
        print(f"  -> 失败: {e}")
        return False

def process_path_pattern(path_pattern, action_func, **kwargs):
    """
    处理路径通配符，批量执行 action_func
    """
    # 去除引号
    path_pattern = path_pattern.strip('"').strip("'")
    
    # 使用 glob 查找所有匹配的文件
    files = glob.glob(path_pattern)
    
    if not files:
        print(f"未找到匹配的文件: {path_pattern}")
        return

    print(f"找到 {len(files)} 个文件，开始批量处理...")
    
    success_count = 0
    for filepath in files:
        if os.path.isfile(filepath):
            if action_func(filepath, **kwargs):
                success_count += 1
    
    print(f"批量处理完成。成功: {success_count}/{len(files)}")

def main():
    while True:
        print("\n" + "="*40)
        print("     DEBUG <-> Supporting 转换工具")
        print("="*40)
        print("1. Supporting (Doc/Txt) -> DEBUG (JSON)")
        print("2. DEBUG (JSON) -> Supporting (Txt/Doc)")
        print("3. Exit")
        print("="*40)

        choice = input("请输入选项 (1-3): ").strip()

        if choice == '1':
            path_pattern = input("请输入 Supporting 文件路径 (支持通配符 *): ").strip()
            process_path_pattern(path_pattern, support_to_debug)

        elif choice == '2':
            path_pattern = input("请输入 DEBUG 文件路径 (支持通配符 *): ").strip()
            
            # 为了批量处理方便，这里统一询问一次格式，应用于所有文件
            # 如果想对每个文件单独问，可以把 input 移到 debug_to_support 内部（但不推荐）
            fmt_choice = input("请输入输出文件格式 (txt/md/docx) [默认 txt]: ").strip().lower()
            if not fmt_choice: fmt_choice = "txt"
            
            process_path_pattern(path_pattern, debug_to_support, output_format=fmt_choice)
        
        elif choice == '3':
            print("再见！")
            break
        
        else:
            print("无效的选项，请重试。")

if __name__ == "__main__":
    main()