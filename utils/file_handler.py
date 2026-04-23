"""
Project: Joey - The Voice-Enabled End-to-End Career Copilot
File: utils/file_handler.py
Description: 
    This module provides a unified I/O interface for file operations.
    
    Key Capabilities:
    1. Format Agnostic Loading: Automatically detects file extensions (.pdf, .docx, .txt, .md)
       and uses the appropriate library to extract text.
    2. Smart Docx Generation: Contains a lightweight Markdown parser to convert 
       generated text (with **bold** syntax) into formatted Word documents.
"""

import os
import logging
from docx import Document
# Imports for Word document formatting control
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
import PyPDF2

class FileLoader:
    """
    A utility class handling all file read/write operations for the Agent.
    """

    def __init__(self):
        self.logger = logging.getLogger("FileLoader")

    def load(self, file_path: str) -> str:
        """
        [Unified Read Interface]
        Reads content from a file, automatically handling different formats based on extension.
        
        Args:
            file_path: Absolute or relative path to the file.
        
        Returns:
            str: Extracted text content, or empty string if failed.
        """
        if not os.path.exists(file_path):
            self.logger.warning(f"File not found: {file_path}")
            return ""
            
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        try:
            # Dispatch to specific handlers
            if ext == '.docx':
                doc = Document(file_path)
                return '\n'.join([p.text for p in doc.paragraphs])
            
            elif ext == '.pdf':
                text = []
                # Binary read mode for PDF
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text.append(page.extract_text() or "")
                return '\n'.join(text)
            
            else: # Default to text/markdown
                # 'errors=ignore' prevents crashing on weird encoding characters
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
        except Exception as e:
            self.logger.error(f"Read error {file_path}: {e}")
            return ""

    def save_docx(self, content: str, file_path: str):
        """
        [Smart Write Interface]
        Saves raw text string into a formatted .docx file.
        
        Feature:
        It implements a mini-parser to translate Markdown-style bold syntax (**text**)
        into actual Word document bold formatting.

        Updated formatting:
        - Font: DengXian (Equally for English and East Asian)
        - Size: 11 pt
        - Alignment: Justified
        - Line Spacing: 1.16
        - Spacing: Before 0pt, After 8pt
        """
        try:
            doc = Document()
            
            # Process line by line to maintain paragraph structure
            lines = content.split('\n')
            
            for line in lines:
                stripped_line = line.strip()
                if not stripped_line:
                    continue # Skip empty lines to avoid huge gaps

                p = doc.add_paragraph()

                # --- 1. 修改段落格式 (Paragraph Format) ---
                # 两侧对齐
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # 行间距 = 1.16倍
                p.paragraph_format.line_spacing = 1.16
                # 段前 0 磅
                p.paragraph_format.space_before = Pt(0)
                # 段后 8 磅
                p.paragraph_format.space_after = Pt(8)

                # --- Markdown Bold Parsing Logic ---
                # The string is split by '**'.
                # Example: "Skill: **Python** expert" -> ['Skill: ', 'Python', ' expert']
                # - Even indices (0, 2) are normal text.
                # - Odd indices (1) are the text inside **...** (to be bolded).
                parts = stripped_line.split('**')
                
                for i, part in enumerate(parts):
                    if not part: continue 
                    
                    run = p.add_run(part)

                    # --- 2. 修改字体格式 (Font Format) ---
                    # 设置字号 11
                    run.font.size = Pt(11)
                    # 设置西文字体为 等线 (DengXian)
                    run.font.name = 'DengXian'
                    # 设置中文字体为 等线 (DengXian)
                    # 这一步对于显示中文至关重要
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'DengXian')
                    
                    # If index is odd, it was surrounded by **, so make it bold.
                    if i % 2 == 1:
                        run.bold = True
                        # Optional: Add color or font size tweaks here if needed

            doc.save(file_path)
            self.logger.info(f"Saved Docx: {file_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to save docx {file_path}: {e}")
            return False
            